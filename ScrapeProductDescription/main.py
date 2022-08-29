import docx
import pprint

doc = docx.Document('Stones.docx')

products = set()

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if len(cell.paragraphs) > 1:
                products.add(cell.paragraphs[1].text.strip())
            else:
                continue


products.remove("")
pprint.pprint(products)